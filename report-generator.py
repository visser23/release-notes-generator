import requests
from docx import Document

# Jira and Confluence details
JIRA_BASE_URL = 'https://servita-uk.atlassian.net/jira'
JIRA_API_USERNAME = 'matt.visser@servita.com'
JIRA_API_TOKEN = 'ATATT3xFfGF0h1WuFZJy9jXFViCDlYVYILzWtqukLDHbEH2_KJIeP2zMx5m4yF2ENPYnUkMfaRRw9pSgDx8bdS3DH_bqAGglcsIF8RetSA4CxR2DSxIvmJXap-xrarRh4lyigKPnF-fyV75L34pyuedQws8Nv9RLXqwePRCz3BjuA_eQ16aDYJ0=5A479038'
CONFLUENCE_BASE_URL = 'https://servita-uk.atlassian.net/wiki'
CONFLUENCE_PAGE_ID = '139853825'
CONFLUENCE_API_USERNAME = 'matt.visser@servita.com'
CONFLUENCE_API_TOKEN = 'ATATT3xFfGF0h1WuFZJy9jXFViCDlYVYILzWtqukLDHbEH2_KJIeP2zMx5m4yF2ENPYnUkMfaRRw9pSgDx8bdS3DH_bqAGglcsIF8RetSA4CxR2DSxIvmJXap-xrarRh4lyigKPnF-fyV75L34pyuedQws8Nv9RLXqwePRCz3BjuA_eQ16aDYJ0=5A479038'

def generate_release_notes(released_fix_version):
    # Retrieve Jira issues for the released fix version
    issues_url = f'{JIRA_BASE_URL}/rest/api/2/search'
    query = f'fixVersion="{released_fix_version}"'
    headers = {'Content-Type': 'application/json'}
    auth = (JIRA_API_USERNAME, JIRA_API_TOKEN)
    payload = {'jql': query, 'fields': 'summary,description'}
    response = requests.get(issues_url, headers=headers, auth=auth, params=payload)
    issues = response.json().get('issues', [])

    # Generate the release notes document
    doc = Document()
    doc.add_heading('Release Notes', level=1)
    for issue in issues:
        issue_key = issue['key']
        summary = issue['fields']['summary']
        description = issue['fields']['description']
        doc.add_heading(issue_key, level=2)
        doc.add_paragraph(f'Summary: {summary}')
        doc.add_paragraph(f'Description: {description}')
        doc.add_paragraph('')  # Add an empty line between issues

    # Save the release notes document
    doc.save('release_notes.docx')

    # Update the Confluence page with the release notes document
    confluence_url = f'{CONFLUENCE_BASE_URL}/wiki/rest/api/content/{CONFLUENCE_PAGE_ID}/child/attachment'
    headers = {'X-Atlassian-Token': 'no-check'}
    auth = (CONFLUENCE_API_USERNAME, CONFLUENCE_API_TOKEN)
    files = {'file': open('release_notes.docx', 'rb')}
    response = requests.post(confluence_url, headers=headers, auth=auth, files=files)

    if response.status_code == 200:
        print('Release notes generated and uploaded successfully.')
    else:
        print('Failed to upload release notes.')

# Example usage
released_fix_version = '1.0.0'  # Replace with the actual released fix version
generate_release_notes(released_fix_version)
